from pathlib import Path
from docx import Document

OUTPUT_DIR = Path("backend/output")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

MASTER_FILE = OUTPUT_DIR / "MasterResume.docx"


def add_list(document, heading, items):

    document.add_heading(heading, level=2)

    if not items:
        document.add_paragraph("N/A")
        return

    for item in items:

        if isinstance(item, dict):

            for key, value in item.items():

                if isinstance(value, list):

                    document.add_paragraph(
                        f"{key.capitalize()}:",
                        style="List Bullet"
                    )

                    for sub in value:
                        document.add_paragraph(
                            str(sub),
                            style="List Bullet 2"
                        )

                else:

                    document.add_paragraph(
                        f"{key.capitalize()}: {value}",
                        style="List Bullet"
                    )

        else:

            document.add_paragraph(
                str(item),
                style="List Bullet"
            )


def append_resume(data):

    if MASTER_FILE.exists():

        document = Document(MASTER_FILE)

        document.add_page_break()

    else:

        document = Document()

    document.add_heading(
        data.get("name", "Unknown Candidate"),
        level=1
    )

    document.add_heading("Contact Information", level=2)

    document.add_paragraph(
        f"Email: {data.get('email','')}"
    )

    document.add_paragraph(
        f"Phone: {data.get('phone','')}"
    )

    document.add_paragraph(
        f"Location: {data.get('location','')}"
    )

    document.add_paragraph(
        f"LinkedIn: {data.get('linkedin','')}"
    )

    document.add_paragraph(
        f"GitHub: {data.get('github','')}"
    )

    document.add_paragraph(
        f"Portfolio: {data.get('portfolio','')}"
    )

    document.add_heading(
        "Professional Summary",
        level=2
    )

    document.add_paragraph(
        data.get("summary", "")
    )

    add_list(
        document,
        "Education",
        data.get("education", [])
    )

    add_list(
        document,
        "Experience",
        data.get("experience", [])
    )

    add_list(
        document,
        "Skills",
        data.get("skills", [])
    )

    add_list(
        document,
        "Projects",
        data.get("projects", [])
    )

    add_list(
        document,
        "Certifications",
        data.get("certifications", [])
    )

    add_list(
        document,
        "Languages",
        data.get("languages", [])
    )

    add_list(
        document,
        "Achievements",
        data.get("achievements", [])
    )

    document.save(MASTER_FILE)

    return str(MASTER_FILE)